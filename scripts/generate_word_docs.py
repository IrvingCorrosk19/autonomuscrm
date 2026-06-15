#!/usr/bin/env python3
"""Generate corporate Word documents from transformed markdown."""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(r"c:\Proyectos\autonomuscrm")
CORP = ROOT / "Documentation" / "Corporate"
WORD = ROOT / "Documentation" / "Word"

VERSION = "2.0.0"
CLASSIFICATION = "Confidencial — Uso interno y clientes autorizados"


def set_segoe(run, size=11, bold=False, color=None):
    run.font.name = "Segoe UI"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Segoe UI")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_header_footer(doc, title: str):
    section = doc.sections[0]
    header = section.header
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.text = "AutonomusCRM"
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_segoe(hp.runs[0], 9, True, RGBColor(0, 120, 212))

    footer = section.footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run(f"{title}  |  v{VERSION}  |  {CLASSIFICATION}  |  Página ")
    set_segoe(run, 8)


def parse_md_to_docx(md_path: Path, docx_path: Path):
    text = md_path.read_text(encoding="utf-8")
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    title = md_path.stem.replace("_", " ")
    add_header_footer(doc, title)

    in_code = False
    in_table = False
    table_rows = []

    for line in text.split("\n"):
        raw = line.rstrip()

        if raw.startswith("```"):
            in_code = not in_code
            if in_code:
                p = doc.add_paragraph()
                r = p.add_run("[Diagrama / bloque técnico — ver versión Markdown]")
                set_segoe(r, 10, False, RGBColor(100, 100, 100))
            continue
        if in_code:
            p = doc.add_paragraph()
            r = p.add_run(raw)
            set_segoe(r, 9)
            p.paragraph_format.left_indent = Inches(0.3)
            continue

        if raw.startswith("|") and "|" in raw[1:]:
            if re.match(r"^\|[-: |]+\|$", raw):
                in_table = True
                continue
            cells = [c.strip() for c in raw.strip("|").split("|")]
            table_rows.append(cells)
            in_table = True
            continue
        elif in_table and table_rows:
            cols = max(len(r) for r in table_rows)
            tbl = doc.add_table(rows=len(table_rows), cols=cols)
            tbl.style = "Table Grid"
            for ri, row in enumerate(table_rows):
                for ci, cell in enumerate(row):
                    if ci < cols:
                        tbl.rows[ri].cells[ci].text = cell
            table_rows = []
            in_table = False

        if not raw:
            continue

        if raw.startswith("# ") and not raw.startswith("## "):
            p = doc.add_heading(raw[2:].strip(), level=0)
            for r in p.runs:
                set_segoe(r, 20, True, RGBColor(0, 120, 212))
        elif raw.startswith("## "):
            p = doc.add_heading(raw[3:].strip(), level=1)
            for r in p.runs:
                set_segoe(r, 16, True)
        elif raw.startswith("### "):
            p = doc.add_heading(raw[4:].strip(), level=2)
            for r in p.runs:
                set_segoe(r, 14, True)
        elif raw.startswith("#### "):
            p = doc.add_heading(raw[5:].strip(), level=3)
            for r in p.runs:
                set_segoe(r, 12, True)
        elif raw.startswith("> **"):
            p = doc.add_paragraph()
            m = re.match(r"> \*\*([A-ZÁÉÍÓÚÑ ]+)\*\*\s*(.*)", raw)
            if m:
                label, rest = m.group(1), m.group(2)
                colors = {
                    "IMPORTANTE": RGBColor(180, 0, 0),
                    "ADVERTENCIA": RGBColor(200, 100, 0),
                    "NOTA": RGBColor(0, 100, 180),
                    "BUENA PRÁCTICA": RGBColor(0, 130, 60),
                    "RIESGO": RGBColor(180, 0, 0),
                    "RECOMENDACIÓN": RGBColor(0, 120, 212),
                }
                r1 = p.add_run(f"{label}: ")
                set_segoe(r1, 11, True, colors.get(label, RGBColor(0, 0, 0)))
                r2 = p.add_run(rest)
                set_segoe(r2, 11)
            else:
                r = p.add_run(raw.lstrip("> "))
                set_segoe(r, 11)
            p.paragraph_format.left_indent = Inches(0.25)
        elif raw.startswith("[CAPTURA:"):
            p = doc.add_paragraph()
            r = p.add_run(raw)
            set_segoe(r, 10, True, RGBColor(128, 128, 128))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # placeholder box
            p2 = doc.add_paragraph()
            r2 = p2.add_run("┌─────────────────────────────────────────┐\n│  Espacio reservado para captura        │\n└─────────────────────────────────────────┘")
            set_segoe(r2, 9, False, RGBColor(180, 180, 180))
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif raw.startswith("**") and raw.endswith("**"):
            p = doc.add_paragraph()
            r = p.add_run(raw.strip("*"))
            set_segoe(r, 11, True)
        else:
            clean = re.sub(r"\*\*([^*]+)\*\*", r"\1", raw)
            clean = re.sub(r"`([^`]+)`", r"\1", clean)
            p = doc.add_paragraph(clean)
            for r in p.runs:
                set_segoe(r, 11)

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))


def main():
    WORD.mkdir(parents=True, exist_ok=True)
    mapping = {
        "Roles/Admin_User_Manual.md": "Admin_User_Manual.docx",
        "Roles/Manager_User_Manual.md": "Manager_User_Manual.docx",
        "Roles/Sales_User_Manual.md": "Sales_User_Manual.docx",
        "Roles/Support_User_Manual.md": "Support_User_Manual.docx",
        "Roles/Viewer_User_Manual.md": "Viewer_User_Manual.docx",
        "ADMIN_OPERATIONS_GUIDE.md": "Admin_Operations_Guide.docx",
        "SALES_PLAYBOOK.md": "Sales_Playbook.docx",
        "SUPPORT_OPERATIONS_GUIDE.md": "Support_Operations_Guide.docx",
        "CUSTOMER_SUCCESS_PLAYBOOK.md": "CustomerSuccess_Playbook.docx",
        "MARKETING_OPERATIONS_GUIDE.md": "Marketing_Operations_Guide.docx",
        "NEW_EMPLOYEE_ONBOARDING.md": "New_Employee_Onboarding.docx",
        "ROLE_DISCOVERY_REPORT.md": "Role_Discovery_Report.docx",
        "ROLE_PERMISSION_MATRIX.md": "Role_Permission_Matrix.docx",
        "SuperAdmin_Status.md": "SuperAdmin_Status.docx",
        "manual-empresarial/02_BUSINESS_FLOWS.md": "Business_Flows.docx",
        "manual-empresarial/08_FAQ.md": "FAQ.docx",
        "manual-empresarial/09_TROUBLESHOOTING.md": "Troubleshooting.docx",
    }
    count = 0
    for rel, docx_name in mapping.items():
        md = CORP / rel
        if not md.exists():
            print(f"SKIP (missing): {rel}")
            continue
        out = WORD / docx_name
        parse_md_to_docx(md, out)
        count += 1
        print(f"OK DOCX: {docx_name}")
    print(f"Generated {count} Word documents -> {WORD}")


if __name__ == "__main__":
    main()
